import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader
from pypdf.errors import PdfReadError

from backend.app.models.document import Document


@dataclass(frozen=True)
class ParsedPage:
    page_number: int
    text: str


@dataclass(frozen=True)
class ParsedHeading:
    level: int
    text: str
    page_number: int = 1


@dataclass(frozen=True)
class ParsedDocument:
    document_id: str
    file_type: str
    pages: list[ParsedPage]
    headings: list[ParsedHeading]

    @property
    def text(self) -> str:
        return "\n\n".join(page.text for page in self.pages if page.text)


class DocumentParsingError(Exception):
    message = "Document parsing failed"


class UnsupportedDocumentTypeError(DocumentParsingError):
    message = "Unsupported document type"


class DocumentFileNotFoundError(DocumentParsingError):
    message = "Document file not found"


class DocumentTextDecodeError(DocumentParsingError):
    message = "Document text could not be decoded"


HEADING_PATTERN = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")


def parse_document(document: Document) -> ParsedDocument:
    path = Path(document.storage_path)
    if not path.is_file():
        raise DocumentFileNotFoundError

    headings: list[ParsedHeading]
    match document.file_type:
        case "pdf":
            pages = parse_pdf(path)
            headings = []
        case "txt":
            pages = parse_txt(path)
            headings = []
        case "docx":
            pages, headings = parse_docx(path)
        case "md" | "markdown":
            pages, headings = parse_markdown(path)
        case _:
            raise UnsupportedDocumentTypeError

    return ParsedDocument(
        document_id=str(document.id),
        file_type=document.file_type,
        pages=pages,
        headings=headings,
    )


def parse_txt(path: Path) -> list[ParsedPage]:
    try:
        text = normalize_text(path.read_text(encoding="utf-8-sig"))
    except UnicodeDecodeError as exc:
        raise DocumentTextDecodeError from exc
    except OSError as exc:
        raise DocumentParsingError from exc

    return [ParsedPage(page_number=1, text=text)]


def parse_pdf(path: Path) -> list[ParsedPage]:
    try:
        reader = PdfReader(path)
        return [
            ParsedPage(page_number=page_number, text=normalize_text(page.extract_text() or ""))
            for page_number, page in enumerate(reader.pages, start=1)
        ]
    except (PdfReadError, OSError, ValueError) as exc:
        raise DocumentParsingError from exc


def parse_docx(path: Path) -> tuple[list[ParsedPage], list[ParsedHeading]]:
    try:
        docx_document = DocxDocument(str(path))
    except Exception as exc:
        raise DocumentParsingError from exc

    lines: list[str] = []
    headings: list[ParsedHeading] = []
    for paragraph in docx_document.paragraphs:
        text = normalize_inline_text(paragraph.text)
        if not text:
            continue

        style_name = paragraph.style.name if paragraph.style is not None else ""
        heading_level = get_docx_heading_level(style_name)
        if heading_level is not None:
            headings.append(ParsedHeading(level=heading_level, text=text))

        lines.append(text)

    text = normalize_text("\n".join(lines))
    return [ParsedPage(page_number=1, text=text)], headings


def parse_markdown(path: Path) -> tuple[list[ParsedPage], list[ParsedHeading]]:
    try:
        raw_text = path.read_text(encoding="utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentTextDecodeError from exc
    except OSError as exc:
        raise DocumentParsingError from exc

    headings: list[ParsedHeading] = []
    output_lines: list[str] = []
    for line in raw_text.splitlines():
        heading_match = HEADING_PATTERN.match(line.strip())
        if heading_match is not None:
            heading_text = normalize_inline_text(heading_match.group(2))
            headings.append(
                ParsedHeading(
                    level=len(heading_match.group(1)),
                    text=heading_text,
                )
            )
            output_lines.append(heading_text)
            continue

        output_lines.append(line)

    text = normalize_text("\n".join(output_lines))
    return [ParsedPage(page_number=1, text=text)], headings


def get_docx_heading_level(style_name: str) -> int | None:
    match = re.fullmatch(r"Heading\s+([1-6])", style_name)
    if match is None:
        return None
    return int(match.group(1))


def normalize_inline_text(text: str) -> str:
    return re.sub(r"[ \t]+", " ", text.replace("\xa0", " ")).strip()


def normalize_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n").replace("\xa0", " ")
    normalized = "\n".join(normalize_inline_text(line) for line in normalized.split("\n"))
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()
