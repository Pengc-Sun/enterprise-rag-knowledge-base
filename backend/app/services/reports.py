import uuid
from base64 import b64encode
from collections.abc import Iterable
from io import BytesIO
from json import dumps
from pathlib import Path

from docx import Document as DocxDocument
from reportlab.lib.pagesizes import LETTER  # type: ignore[import-untyped]
from reportlab.lib.units import inch  # type: ignore[import-untyped]
from reportlab.pdfgen.canvas import Canvas  # type: ignore[import-untyped]
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from backend.app.models.analysis import AnalysisResult, AnalysisResultStatus, AnalysisTask
from backend.app.models.report import (
    ExportFormat,
    ExportJob,
    ExportJobStatus,
    Report,
    ReportSection,
    ReportSectionStatus,
    ReportStatus,
)
from backend.app.schemas.report import (
    ReportCreate,
    ReportExportCreate,
    ReportPreviewRead,
    ReportSectionCreate,
    ReportSectionGenerateRequest,
    ReportSectionReorderRequest,
    ReportSectionUpdate,
    ReportUpdate,
)

REPORTABLE_ANALYSIS_RESULT_STATUSES = frozenset(
    {
        AnalysisResultStatus.APPROVED.value,
        AnalysisResultStatus.EDITED.value,
    }
)


class ReportSectionSourceError(ValueError):
    pass


class ReportSectionGenerationError(ReportSectionSourceError):
    pass


class ReportSectionOrderingError(ValueError):
    pass


class ReportExportError(ValueError):
    pass


async def list_reports_for_workspace(
    session: AsyncSession,
    workspace_id: uuid.UUID,
) -> list[Report]:
    result = await session.execute(
        select(Report)
        .where(Report.workspace_id == workspace_id)
        .order_by(Report.created_at.desc(), Report.id.desc())
    )
    return list(result.scalars().all())


async def get_report_for_workspace(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    report_id: uuid.UUID,
) -> Report | None:
    result = await session.execute(
        select(Report).where(
            Report.id == report_id,
            Report.workspace_id == workspace_id,
        )
    )
    return result.scalar_one_or_none()


async def create_report(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    created_by: uuid.UUID,
    report_create: ReportCreate,
) -> Report:
    report = Report(
        workspace_id=workspace_id,
        title=report_create.title,
        status=ReportStatus.DRAFT.value,
        created_by=created_by,
    )
    session.add(report)
    await session.commit()
    await session.refresh(report)
    return report


async def update_report(
    session: AsyncSession,
    report: Report,
    report_update: ReportUpdate,
) -> Report:
    update_data = report_update.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(report, field, value)

    await session.commit()
    await session.refresh(report)
    return report


async def list_report_sections_for_report(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    report_id: uuid.UUID,
) -> list[ReportSection]:
    result = await session.execute(
        select(ReportSection)
        .where(
            ReportSection.workspace_id == workspace_id,
            ReportSection.report_id == report_id,
        )
        .order_by(ReportSection.sort_order.asc(), ReportSection.created_at.asc())
    )
    return list(result.scalars().all())


async def get_report_section_for_report(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    report_id: uuid.UUID,
    section_id: uuid.UUID,
) -> ReportSection | None:
    result = await session.execute(
        select(ReportSection).where(
            ReportSection.id == section_id,
            ReportSection.workspace_id == workspace_id,
            ReportSection.report_id == report_id,
        )
    )
    return result.scalar_one_or_none()


async def build_report_preview(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    report: Report,
) -> ReportPreviewRead:
    sections = await list_report_sections_for_report(session, workspace_id, report.id)
    return build_report_preview_from_sections(workspace_id, report, sections)


def build_report_preview_from_sections(
    workspace_id: uuid.UUID,
    report: Report,
    sections: Iterable[ReportSection],
) -> ReportPreviewRead:
    section_list = list(sections)
    return ReportPreviewRead(
        report_id=report.id,
        workspace_id=workspace_id,
        title=report.title,
        status=ReportStatus(report.status),
        section_count=len(section_list),
        markdown=render_report_preview_markdown(report, section_list),
    )


async def get_export_job_for_workspace(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    export_id: uuid.UUID,
) -> ExportJob | None:
    result = await session.execute(
        select(ExportJob).where(
            ExportJob.id == export_id,
            ExportJob.workspace_id == workspace_id,
        )
    )
    return result.scalar_one_or_none()


async def create_report_export(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    report: Report,
    created_by: uuid.UUID,
    export_create: ReportExportCreate,
    export_dir: str = "storage/exports",
) -> ExportJob:
    if export_create.format not in {
        ExportFormat.MARKDOWN,
        ExportFormat.DOCX,
        ExportFormat.PDF,
    }:
        raise ReportExportError("Only markdown, docx, and pdf export are supported")

    sections = await list_report_sections_for_report(session, workspace_id, report.id)
    await validate_report_export_sections(session, workspace_id, sections)
    preview = build_report_preview_from_sections(workspace_id, report, sections)
    export_metadata, export_bytes = build_export_artifact(export_create.format, preview)
    export_id = uuid.uuid4()
    file_path = write_export_file(
        export_dir,
        workspace_id,
        export_id,
        str(export_metadata["filename"]),
        export_bytes,
    )
    export_job = ExportJob(
        id=export_id,
        workspace_id=workspace_id,
        report_id=report.id,
        format=export_create.format.value,
        status=ExportJobStatus.COMPLETED.value,
        file_path=file_path,
        error_message=None,
        created_by=created_by,
        export_metadata=export_metadata,
    )
    report.status = ReportStatus.EXPORTED.value
    session.add(export_job)
    await session.commit()
    await session.refresh(export_job)
    await session.refresh(report)
    return export_job


def build_export_metadata(
    export_format: ExportFormat,
    preview: ReportPreviewRead,
) -> dict[str, object]:
    metadata, _ = build_export_artifact(export_format, preview)
    return metadata


def build_export_artifact(
    export_format: ExportFormat,
    preview: ReportPreviewRead,
) -> tuple[dict[str, object], bytes]:
    metadata: dict[str, object] = {
        "title": preview.title,
        "section_count": preview.section_count,
    }
    if export_format == ExportFormat.MARKDOWN:
        markdown_bytes = preview.markdown.encode("utf-8")
        metadata["markdown"] = preview.markdown
        metadata["content_type"] = "text/markdown; charset=utf-8"
        metadata["filename"] = f"{_export_filename_stem(preview.title)}.md"
        return metadata, markdown_bytes
    if export_format == ExportFormat.DOCX:
        docx_bytes = render_report_docx(preview)
        metadata["docx_base64"] = b64encode(docx_bytes).decode("ascii")
        metadata["content_type"] = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        metadata["filename"] = f"{_export_filename_stem(preview.title)}.docx"
        return metadata, docx_bytes
    if export_format == ExportFormat.PDF:
        pdf_bytes = render_report_pdf(preview)
        metadata["pdf_base64"] = b64encode(pdf_bytes).decode("ascii")
        metadata["content_type"] = "application/pdf"
        metadata["filename"] = f"{_export_filename_stem(preview.title)}.pdf"
        return metadata, pdf_bytes
    raise ReportExportError("Unsupported export format")


def write_export_file(
    export_dir: str,
    workspace_id: uuid.UUID,
    export_id: uuid.UUID,
    filename: str,
    content: bytes,
) -> str:
    export_directory = Path(export_dir) / str(workspace_id) / str(export_id)
    export_directory.mkdir(parents=True, exist_ok=True)
    export_path = export_directory / filename
    export_path.write_bytes(content)
    return export_path.as_posix()


async def validate_report_export_sections(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    sections: Iterable[ReportSection],
) -> None:
    source_result_ids = [
        source_result_id
        for section in sections
        for source_result_id in section.source_result_ids
    ]
    try:
        await validate_report_section_source_results(session, workspace_id, source_result_ids)
    except ReportSectionSourceError as exc:
        raise ReportExportError(
            "Report exports can only include approved or edited analysis results"
        ) from exc


def render_report_docx(preview: ReportPreviewRead) -> bytes:
    document = DocxDocument()
    document.add_heading(preview.title, level=1)

    for line in preview.markdown.splitlines()[1:]:
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("## "):
            document.add_heading(stripped.removeprefix("## ").strip(), level=2)
        elif stripped.startswith("### "):
            document.add_heading(stripped.removeprefix("### ").strip(), level=3)
        elif stripped.startswith("- "):
            document.add_paragraph(stripped.removeprefix("- ").strip(), style="List Bullet")
        else:
            document.add_paragraph(stripped)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def render_report_pdf(preview: ReportPreviewRead) -> bytes:
    output = BytesIO()
    canvas = Canvas(output, pagesize=LETTER)
    page_width, page_height = LETTER
    margin = 0.75 * inch
    max_width = page_width - (2 * margin)
    y = page_height - margin

    def draw_wrapped(
        text: str,
        *,
        font_name: str,
        font_size: int,
        line_height: int,
        indent: float = 0,
    ) -> None:
        nonlocal y
        canvas.setFont(font_name, font_size)
        for wrapped_line in _wrap_pdf_text(canvas, text, max_width - indent, font_name, font_size):
            if y < margin:
                canvas.showPage()
                canvas.setFont(font_name, font_size)
                y = page_height - margin
            canvas.drawString(margin + indent, y, wrapped_line)
            y -= line_height

    draw_wrapped(preview.title, font_name="Helvetica-Bold", font_size=18, line_height=24)
    y -= 8

    for line in preview.markdown.splitlines()[1:]:
        stripped = line.strip()
        if not stripped:
            y -= 8
            continue
        if stripped.startswith("## "):
            y -= 4
            draw_wrapped(
                stripped.removeprefix("## ").strip(),
                font_name="Helvetica-Bold",
                font_size=14,
                line_height=19,
            )
        elif stripped.startswith("### "):
            draw_wrapped(
                stripped.removeprefix("### ").strip(),
                font_name="Helvetica-Bold",
                font_size=12,
                line_height=17,
            )
        elif stripped.startswith("- "):
            draw_wrapped(
                f"- {stripped.removeprefix('- ').strip()}",
                font_name="Helvetica",
                font_size=11,
                line_height=15,
                indent=14,
            )
        else:
            draw_wrapped(stripped, font_name="Helvetica", font_size=11, line_height=15)

    canvas.save()
    return output.getvalue()


async def create_report_section(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    report_id: uuid.UUID,
    section_create: ReportSectionCreate,
) -> ReportSection:
    source_result_ids = await validate_report_section_source_results(
        session,
        workspace_id,
        section_create.source_result_ids,
    )
    section = ReportSection(
        workspace_id=workspace_id,
        report_id=report_id,
        template_section_key=section_create.template_section_key,
        title=section_create.title,
        body_markdown=section_create.body_markdown,
        source_task_keys=section_create.source_task_keys,
        source_result_ids=source_result_ids,
        sort_order=section_create.sort_order,
        status=ReportSectionStatus.DRAFT.value,
    )
    session.add(section)
    await session.commit()
    await session.refresh(section)
    return section


async def update_report_section(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    section: ReportSection,
    section_update: ReportSectionUpdate,
) -> ReportSection:
    update_data = section_update.model_dump(exclude_unset=True)
    if "source_result_ids" in update_data:
        source_result_ids = update_data["source_result_ids"] or []
        update_data["source_result_ids"] = await validate_report_section_source_results(
            session,
            workspace_id,
            source_result_ids,
        )

    for field, value in update_data.items():
        setattr(section, field, value)

    await session.commit()
    await session.refresh(section)
    return section


async def reorder_report_sections(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    report_id: uuid.UUID,
    reorder_request: ReportSectionReorderRequest,
) -> list[ReportSection]:
    section_ids = [item.section_id for item in reorder_request.sections]
    if len(section_ids) != len(set(section_ids)):
        raise ReportSectionOrderingError("Report section ordering cannot contain duplicates")

    result = await session.execute(
        select(ReportSection).where(
            ReportSection.workspace_id == workspace_id,
            ReportSection.report_id == report_id,
            ReportSection.id.in_(section_ids),
        )
    )
    sections = list(result.scalars().all())
    sections_by_id = {section.id: section for section in sections}
    missing_section_ids = [
        section_id for section_id in section_ids if section_id not in sections_by_id
    ]
    if missing_section_ids:
        raise ReportSectionOrderingError(
            "Report section ordering contains sections outside this report"
        )

    for item in reorder_request.sections:
        sections_by_id[item.section_id].sort_order = item.sort_order

    await session.commit()
    for section in sections:
        await session.refresh(section)

    return sorted(
        sections,
        key=lambda section: (section.sort_order, section.created_at, str(section.id)),
    )


def render_report_preview_markdown(
    report: Report,
    sections: Iterable[ReportSection],
) -> str:
    lines = [f"# {_markdown_heading_text(report.title)}", ""]
    section_list = list(sections)
    if not section_list:
        lines.append("_No report sections yet._")
        return "\n".join(lines).strip() + "\n"

    for section in section_list:
        lines.extend([f"## {_markdown_heading_text(section.title)}", ""])
        body_markdown = section.body_markdown.strip()
        lines.append(body_markdown or "_No content yet._")
        lines.append("")

    return "\n".join(lines).strip() + "\n"


async def generate_report_section_from_results(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    report_id: uuid.UUID,
    generation: ReportSectionGenerateRequest,
) -> ReportSection:
    result_ids = _dedupe_uuids(generation.analysis_result_ids)
    rows = await _get_reportable_analysis_result_rows(session, workspace_id, result_ids)
    rows_by_id = {analysis_result.id: (analysis_result, task) for analysis_result, task in rows}

    missing_result_ids = [result_id for result_id in result_ids if result_id not in rows_by_id]
    if missing_result_ids:
        raise ReportSectionGenerationError(
            "Analysis results must exist in this workspace and be approved or edited"
        )

    ordered_rows = [rows_by_id[result_id] for result_id in result_ids]
    section = ReportSection(
        workspace_id=workspace_id,
        report_id=report_id,
        template_section_key=generation.template_section_key,
        title=generation.title or _build_generated_section_title(ordered_rows),
        body_markdown=_build_generated_section_markdown(ordered_rows),
        source_task_keys=_collect_source_task_keys(task for _, task in ordered_rows),
        source_result_ids=[str(analysis_result.id) for analysis_result, _ in ordered_rows],
        sort_order=generation.sort_order,
        status=ReportSectionStatus.DRAFT.value,
    )
    session.add(section)
    await session.commit()
    await session.refresh(section)
    return section


async def validate_report_section_source_results(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    source_result_ids: Iterable[str],
) -> list[str]:
    result_ids = _parse_source_result_ids(source_result_ids)
    if not result_ids:
        return []

    rows = await _get_reportable_analysis_result_rows(session, workspace_id, result_ids)
    reportable_result_ids = {analysis_result.id for analysis_result, _ in rows}
    missing_result_ids = [
        result_id for result_id in result_ids if result_id not in reportable_result_ids
    ]
    if missing_result_ids:
        raise ReportSectionSourceError(
            "Report section source results must exist in this workspace and be approved or edited"
        )
    return [str(result_id) for result_id in result_ids]


async def _get_reportable_analysis_result_rows(
    session: AsyncSession,
    workspace_id: uuid.UUID,
    result_ids: list[uuid.UUID],
) -> list[tuple[AnalysisResult, AnalysisTask]]:
    if not result_ids:
        return []

    result = await session.execute(
        select(AnalysisResult, AnalysisTask)
        .join(AnalysisTask, AnalysisResult.analysis_task_id == AnalysisTask.id)
        .where(
            AnalysisResult.workspace_id == workspace_id,
            AnalysisResult.id.in_(result_ids),
            AnalysisResult.status.in_(REPORTABLE_ANALYSIS_RESULT_STATUSES),
        )
    )
    return [(analysis_result, task) for analysis_result, task in result.all()]


def _parse_source_result_ids(source_result_ids: Iterable[str]) -> list[uuid.UUID]:
    parsed_ids: list[uuid.UUID] = []
    for source_result_id in source_result_ids:
        try:
            parsed_ids.append(uuid.UUID(source_result_id))
        except ValueError as exc:
            raise ReportSectionSourceError(
                "Report section source results must be valid UUIDs"
            ) from exc
    return _dedupe_uuids(parsed_ids)


def _dedupe_uuids(values: Iterable[uuid.UUID]) -> list[uuid.UUID]:
    seen: set[uuid.UUID] = set()
    deduped: list[uuid.UUID] = []
    for value in values:
        if value not in seen:
            deduped.append(value)
            seen.add(value)
    return deduped


def _build_generated_section_title(
    rows: list[tuple[AnalysisResult, AnalysisTask]],
) -> str:
    if len(rows) == 1:
        return rows[0][1].name
    return "Draft Report Section"


def _build_generated_section_markdown(
    rows: list[tuple[AnalysisResult, AnalysisTask]],
) -> str:
    blocks: list[str] = []
    for index, (analysis_result, task) in enumerate(rows, start=1):
        heading = task.name if len(rows) == 1 else f"{index}. {task.name}"
        blocks.append(f"### {heading}")
        blocks.append("")
        blocks.append(_json_block(analysis_result.result))

        if analysis_result.citations:
            blocks.append("")
            blocks.append("#### Citations")
            blocks.extend(
                f"- {_format_citation(citation)}" for citation in analysis_result.citations
            )

        blocks.append("")

    return "\n".join(blocks).strip()


def _json_block(value: object) -> str:
    return f"```json\n{dumps(value, ensure_ascii=False, indent=2, sort_keys=True)}\n```"


def _format_citation(citation: dict[str, object]) -> str:
    title = str(
        citation.get("document_title")
        or citation.get("document_name")
        or citation.get("file_name")
        or citation.get("source")
        or "Source"
    )
    page = citation.get("page")
    chunk_id = citation.get("chunk_id")
    suffix_parts = []
    if page is not None:
        suffix_parts.append(f"page {page}")
    if chunk_id is not None:
        suffix_parts.append(f"chunk {chunk_id}")
    if not suffix_parts:
        return title
    return f"{title} ({', '.join(suffix_parts)})"


def _collect_source_task_keys(tasks: Iterable[AnalysisTask]) -> list[str]:
    keys: list[str] = []
    seen: set[str] = set()
    for task in tasks:
        if task.template_task_key and task.template_task_key not in seen:
            keys.append(task.template_task_key)
            seen.add(task.template_task_key)
    return keys


def _markdown_heading_text(value: str) -> str:
    return " ".join(value.split()).strip() or "Untitled"


def _export_filename_stem(value: str) -> str:
    slug = "".join(character.lower() if character.isalnum() else "-" for character in value)
    compact_slug = "-".join(part for part in slug.split("-") if part)
    return compact_slug[:120] or "report"


def _wrap_pdf_text(
    canvas: Canvas,
    text: str,
    max_width: float,
    font_name: str,
    font_size: int,
) -> list[str]:
    words = text.split()
    if not words:
        return [""]

    lines: list[str] = []
    current_line = words[0]
    for word in words[1:]:
        candidate = f"{current_line} {word}"
        if canvas.stringWidth(candidate, font_name, font_size) <= max_width:
            current_line = candidate
        else:
            lines.append(current_line)
            current_line = word

    lines.append(current_line)
    return lines
