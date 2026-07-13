import uuid
from datetime import UTC, datetime
from pathlib import Path

import pytest
from docx import Document as DocxDocument

from backend.app.models.document import Document
from backend.app.services.document_parsers import (
    DocumentFileNotFoundError,
    DocumentParsingError,
    DocumentTextDecodeError,
    UnsupportedDocumentTypeError,
    parse_document,
)


def make_document(path: Path, file_type: str) -> Document:
    now = datetime.now(UTC)
    return Document(
        id=uuid.uuid4(),
        knowledge_base_id=uuid.uuid4(),
        filename=path.name,
        file_type=file_type,
        file_size=path.stat().st_size if path.exists() else 0,
        file_hash="a" * 64,
        storage_path=path.as_posix(),
        status="uploaded",
        error_message=None,
        created_by=uuid.uuid4(),
        created_at=now,
        updated_at=now,
    )


def make_pdf_bytes() -> bytes:
    first_stream = b"BT /F1 24 Tf 100 700 Td (First page text) Tj ET"
    second_stream = b"BT /F1 24 Tf 100 700 Td (Second page text) Tj ET"
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R 6 0 R] /Count 2 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(first_stream), first_stream),
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 4 0 R >> >> /Contents 7 0 R >>"
        ),
        b"<< /Length %d >>\nstream\n%s\nendstream" % (len(second_stream), second_stream),
    ]

    output = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for object_number, object_body in enumerate(objects, start=1):
        offsets.append(len(output))
        output.extend(f"{object_number} 0 obj\n".encode())
        output.extend(object_body)
        output.extend(b"\nendobj\n")

    xref_offset = len(output)
    output.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    output.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        output.extend(f"{offset:010d} 00000 n \n".encode())
    output.extend(
        f"trailer\n<< /Root 1 0 R /Size {len(objects) + 1} >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode()
    )
    return bytes(output)


def test_parse_txt_document_returns_single_page(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("hello\nworld", encoding="utf-8")
    document = make_document(path, "txt")

    parsed_document = parse_document(document)

    assert parsed_document.document_id == str(document.id)
    assert parsed_document.file_type == "txt"
    assert len(parsed_document.pages) == 1
    assert parsed_document.pages[0].page_number == 1
    assert parsed_document.pages[0].text == "hello\nworld"
    assert parsed_document.text == "hello\nworld"


def test_parse_txt_document_strips_bom_and_normalizes_nonbreaking_spaces(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("\ufeffHello\u00a0\u00a0world\n\n\nNext   line", encoding="utf-8")
    document = make_document(path, "txt")

    parsed_document = parse_document(document)

    assert parsed_document.pages[0].text == "Hello world\n\nNext line"


def test_parse_pdf_document_preserves_page_numbers(tmp_path: Path) -> None:
    path = tmp_path / "handbook.pdf"
    path.write_bytes(make_pdf_bytes())
    document = make_document(path, "pdf")

    parsed_document = parse_document(document)

    assert [page.page_number for page in parsed_document.pages] == [1, 2]
    assert [page.text for page in parsed_document.pages] == [
        "First page text",
        "Second page text",
    ]
    assert parsed_document.text == "First page text\n\nSecond page text"


def test_parse_document_rejects_missing_file(tmp_path: Path) -> None:
    document = make_document(tmp_path / "missing.txt", "txt")

    with pytest.raises(DocumentFileNotFoundError):
        parse_document(document)


def test_parse_markdown_document_extracts_headings_and_normalizes_text(tmp_path: Path) -> None:
    path = tmp_path / "notes.md"
    path.write_text(
        "# Product Plan\r\n\r\nSome   text\n\n\n## Risks ##\n- Item",
        encoding="utf-8",
    )
    document = make_document(path, "md")

    parsed_document = parse_document(document)

    assert parsed_document.pages[0].page_number == 1
    assert parsed_document.pages[0].text == "Product Plan\n\nSome text\n\nRisks\n- Item"
    assert [(heading.level, heading.text) for heading in parsed_document.headings] == [
        (1, "Product Plan"),
        (2, "Risks"),
    ]


def test_parse_markdown_alias_document(tmp_path: Path) -> None:
    path = tmp_path / "notes.markdown"
    path.write_text("# Handbook", encoding="utf-8")
    document = make_document(path, "markdown")

    parsed_document = parse_document(document)

    assert parsed_document.file_type == "markdown"
    assert parsed_document.pages[0].text == "Handbook"
    assert parsed_document.headings[0].text == "Handbook"


def test_parse_docx_document_extracts_headings_and_normalizes_text(tmp_path: Path) -> None:
    path = tmp_path / "handbook.docx"
    docx_document = DocxDocument()
    docx_document.add_heading("Engineering Handbook", level=1)
    docx_document.add_paragraph("First   paragraph")
    docx_document.add_heading("Operations", level=2)
    docx_document.add_paragraph("Second paragraph")
    docx_document.save(str(path))
    document = make_document(path, "docx")

    parsed_document = parse_document(document)

    assert parsed_document.pages[0].page_number == 1
    assert parsed_document.pages[0].text == (
        "Engineering Handbook\nFirst paragraph\nOperations\nSecond paragraph"
    )
    assert [(heading.level, heading.text) for heading in parsed_document.headings] == [
        (1, "Engineering Handbook"),
        (2, "Operations"),
    ]


def test_parse_document_rejects_unsupported_file_type(tmp_path: Path) -> None:
    path = tmp_path / "notes.html"
    path.write_text("<h1>Heading</h1>", encoding="utf-8")
    document = make_document(path, "html")

    with pytest.raises(UnsupportedDocumentTypeError):
        parse_document(document)


def test_parse_txt_document_wraps_decode_errors(tmp_path: Path) -> None:
    path = tmp_path / "bad.txt"
    path.write_bytes(b"\xff\xfe\xfa")
    document = make_document(path, "txt")

    with pytest.raises(DocumentTextDecodeError):
        parse_document(document)


def test_parse_pdf_document_wraps_parser_errors(tmp_path: Path) -> None:
    path = tmp_path / "bad.pdf"
    path.write_bytes(b"not a pdf")
    document = make_document(path, "pdf")

    with pytest.raises(DocumentParsingError):
        parse_document(document)


def test_parse_docx_document_wraps_parser_errors(tmp_path: Path) -> None:
    path = tmp_path / "bad.docx"
    path.write_bytes(b"not a docx")
    document = make_document(path, "docx")

    with pytest.raises(DocumentParsingError):
        parse_document(document)
